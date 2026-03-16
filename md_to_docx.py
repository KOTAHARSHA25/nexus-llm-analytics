import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys
import os

def clean_text(text):
    # Remove markdown bold/italic markers for docx insertion (simplified)
    # A full parser would handle this per-run, but for now we want clean text
    # or we can use a regex to handle bolding within the add_paragraph logic
    return text

def parse_markdown_table(table_block):
    lines = table_block.strip().split('\n')
    if len(lines) < 2:
        return None
    
    # Check for separator line
    if not re.match(r'^[|\s\-:]+$', lines[1]):
        return None

    headers = [c.strip() for c in lines[0].strip('|').split('|')]
    rows = []
    for line in lines[2:]:
        row = [c.strip() for c in line.strip('|').split('|')]
        if len(row) == len(headers):
            rows.append(row)
            
    return {'headers': headers, 'rows': rows}

def add_markdown_formatted_text(paragraph, text):
    # Simple bold/italic parser helper
    # Matches **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()
    
    # Set style basics if needed
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line[level:].strip()
            doc.add_heading(text, level=min(level, 9))
            
        # Tables (look ahead for table structure)
        elif line.startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1 # Step back as loop increments
            
            table_data = parse_markdown_table('\n'.join(table_lines))
            if table_data:
                table = doc.add_table(rows=1, cols=len(table_data['headers']))
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(table_data['headers']):
                    hdr_cells[j].text = header
                    for run in hdr_cells[j].paragraphs[0].runs:
                        run.bold = True
                        
                # Data rows
                for row_data in table_data['rows']:
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(row_data):
                        row_cells[j].text = cell_text
            else:
                # If parse failed, treat as text
                p = doc.add_paragraph()
                add_markdown_formatted_text(p, line)

        # Blockquotes
        elif line.startswith('>'):
            p = doc.add_paragraph()
            add_markdown_formatted_text(p, line[1:].strip())
            p.style = 'Quote'
            
        # Standard Paragraphs (skip empty lines)
        elif line:
            p = doc.add_paragraph()
            add_markdown_formatted_text(p, line)
            
        i += 1

    doc.save(docx_path)
    return f"Successfully converted {md_path} to {docx_path}"

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input_md> <output_docx>")
        sys.exit(1)
    
    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    print(md_to_docx(md_path, docx_path))
