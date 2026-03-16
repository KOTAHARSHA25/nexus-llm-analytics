from fpdf import FPDF
from docx import Document
import random
from pathlib import Path

DATA_DIR = Path("data/stress_test_docs")
DATA_DIR.mkdir(parents=True, exist_ok=True)

def generate_massive_text():
    """Generates a 50KB text file with a specific hidden key."""
    content = "This is a standard clause in the contract. " * 2000
    # The Needle
    content += "\n\nCRITICAL CLAUSE: The release code for the project is 77-ALPHA-OMEGA.\n\n"
    content += "End of document filler. " * 1000
    
    with open(DATA_DIR / "massive_contract.txt", "w") as f:
        f.write(content)
    print("Generated massive_contract.txt")

def generate_complex_docx():
    """Generates a DOCX with headers and tables."""
    doc = Document()
    doc.add_heading('Financial Report 2025', 0)
    
    doc.add_paragraph('This is a confidential analysis of the Q4 performance.')
    
    # Table with data to extract
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Department'
    hdr_cells[1].text = 'Revenue'
    hdr_cells[2].text = 'Status'
    
    row = table.add_row().cells
    row[0].text = 'R&D'
    row[1].text = '$1,500,000'
    row[2].text = 'Active' # Key info
    
    doc.save(DATA_DIR / "complex_report.docx")
    print("Generated complex_report.docx")

def generate_pdf():
    """Generates a PDF with standard text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Project Nexus Final Audit", ln=1, align="C")
    pdf.cell(200, 10, txt="Approved by: Director Johnson", ln=2, align="L")
    pdf.output(DATA_DIR / "audit_final.pdf")
    print("Generated audit_final.pdf")

if __name__ == "__main__":
    generate_massive_text()
    try:
        generate_complex_docx()
    except Exception as e:
        print(f"Skipping DOCX: {e}")
    try:
        generate_pdf()
    except Exception as e:
        print(f"Skipping PDF: {e}")
