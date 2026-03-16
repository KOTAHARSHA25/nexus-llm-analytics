from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx():
    doc = Document()
    
    # Title
    heading = doc.add_heading('Response to Reviewers', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    p = doc.add_paragraph()
    run = p.add_run('Paper Title: ')
    run.bold = True
    p.add_run('Domain Flexible Autonomous Data Analysis Agent Leveraging Large Language Models\n')
    run = p.add_run('Submission Target: ')
    run.bold = True
    p.add_run('IEEE')

    doc.add_paragraph('---')

    # Comment 1
    doc.add_heading('Reviewer Comment 1:', level=2)
    p = doc.add_paragraph()
    run = p.add_run('The author should improve the comparison with recent agentic and tool-augmented LLM frameworks.')
    run.italic = True
    
    doc.add_heading('Response:', level=3)
    p = doc.add_paragraph()
    p.add_run('We have significantly enhanced the comparative analysis in the revised manuscript. Specifically, we added ')
    p.add_run('Section II.E (Summary of Differences)').bold = True
    p.add_run(' and ')
    p.add_run('Table I').bold = True
    p.add_run(', which provides a side-by-side comparison of our ')
    p.add_run('DF-Agent').bold = True
    p.add_run(' against leading frameworks such as ')
    p.add_run('AutoGen').bold = True
    p.add_run(', ')
    p.add_run('AgentVerse').bold = True
    p.add_run(', and ')
    p.add_run('Mergen').bold = True
    p.add_run('.')

    p = doc.add_paragraph()
    p.add_run('Changes Made:').bold = True
    doc.add_paragraph('• Table I explicitly contrasts features including Architecture, Routing Logic, Execution Environment, Error Recovery capabilities, and Latency.', style='List Bullet')
    doc.add_paragraph('• Section II.A & II.B now include expanded discussions on the limitations of "Conversable Agents" (e.g., lack of collaborative efficiency) and the specific gaps in scientific code generation that our sandbox approach addresses.', style='List Bullet')

    doc.add_paragraph('---')

    # Comment 2
    doc.add_heading('Reviewer Comment 2:', level=2)
    p = doc.add_paragraph()
    run = p.add_run('What is the impact of different orchestration heuristics on system latency and accuracy, should be clarified.')
    run.italic = True

    doc.add_heading('Response:', level=3)
    p = doc.add_paragraph()
    p.add_run('We have clarified the impact of our orchestration heuristics by including quantitative performance metrics in ')
    p.add_run('Section VI (Results)').bold = True
    p.add_run('. We now present a detailed breakdown of latency and success rates across different query complexity levels (Simple, Moderate, Complex).')

    p = doc.add_paragraph()
    p.add_run('Changes Made:').bold = True
    doc.add_paragraph('• Figure 4 (Module Efficiency Quadrant Analysis) and Figure 5 (Performance Trade-off) have been added to visually quantify these impacts.', style='List Bullet')
    p = doc.add_paragraph('• Section VI text explicitly states:', style='List Bullet')
    
    sub = doc.add_paragraph(style='List Bullet 2')
    sub.add_run('Simple queries: 100% success rate with ~187.8s latency.')
    sub = doc.add_paragraph(style='List Bullet 2')
    sub.add_run('Moderate queries: 71.4% success rate with ~355.1s latency.')
    sub = doc.add_paragraph(style='List Bullet 2')
    sub.add_run('Complex queries: 55.5% success rate with ~383.8s latency.')
    
    doc.add_paragraph('This data demonstrates how our semantic router optimizes for efficiency by evading heavy computations for simpler tasks.', style='List Bullet')

    doc.add_paragraph('---')

    # Comment 3
    doc.add_heading('Reviewer Comment 3:', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Kindly clarify how error correction mechanisms are evaluated under diverse failure scenarios.')
    run.italic = True

    doc.add_heading('Response:', level=3)
    p = doc.add_paragraph()
    p.add_run('We have expanded the description of our error correction evaluation. The revised manuscript details the performance of the ')
    p.add_run('Self-Correction (CoT)').bold = True
    p.add_run(' module and the "Generator-Critic" loop.')

    p = doc.add_paragraph()
    p.add_run('Changes Made:').bold = True
    doc.add_paragraph('• Section IV.C (Sandbox Execution and Validation) details the "Two Friend" loop where the Critic agent autonomously requests corrections upon detecting syntax errors or execution failures (up to 3 retries).', style='List Bullet')
    doc.add_paragraph('• Section VI reports that the Self-Correction module has the highest error rate (35%) but contributes to a 70% autonomous error recovery rate (mentioned in the Abstract and Section I). We identify that this mechanism is crucial for recovering from grammatical and syntax errors prior to user involvement.', style='List Bullet')

    doc.add_paragraph('---')

    # Comment 4
    doc.add_heading('Reviewer Comment 4:', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Can learning-based routing improve agent selection and execution efficiency, should be clarified.')
    run.italic = True

    doc.add_heading('Response:', level=3)
    p = doc.add_paragraph()
    p.add_run('We have addressed the potential of learning-based routing as a key area for future improvement, contrasting it with our current heuristic approach.')

    p = doc.add_paragraph()
    p.add_run('Changes Made:').bold = True
    doc.add_paragraph('• Section VII (Conclusion & Future Directions) explicitly proposes replacing static heuristic routing with machine learning classifiers based on user interaction logs.', style='List Bullet')
    doc.add_paragraph('• We project that this shift to learning-based routing will result in a 5% to 10% improvement in routing accuracy/precision compared to the current static methods.', style='List Bullet')

    doc.add_paragraph('---')

    # Comment 5
    doc.add_heading('Reviewer Comment 5:', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Author should address the absence of standardized datasets for evaluating domain-flexible analytical performance.')
    run.italic = True

    doc.add_heading('Response:', level=3)
    p = doc.add_paragraph()
    p.add_run('We acknowledge the lack of standardized datasets for this specific domain-flexible autonomous analysis task. In the revised paper, we clarify our validation approach using diverse real-world datasets to ensure domain transferability.')

    p = doc.add_paragraph()
    p.add_run('Changes Made:').bold = True
    doc.add_paragraph('• Abstract and Section I emphasize that our validation focuses on "financial, IoT and unstructured data" to demonstrate the "broad epicentre of domain transferability" (Stability Score: 0.92).', style='List Bullet')
    doc.add_paragraph('• We have framed the evaluation to prioritize real-world applicability across disparate domains rather than single-domain benchmarks, addressing the "trilemma of autonomous analytic" (Executability, Domain Flexibility, Reliability) described in Section I.', style='List Bullet')

    # Save
    base_dir = r"c:\Users\mitta\.gemini\antigravity\brain\267ee7b6-76db-4cfc-9ad1-d69c31f04cda"
    output_path = os.path.join(base_dir, "Response_to_Reviewers.docx")
    doc.save(output_path)
    print(f"Docx saved to: {output_path}")

if __name__ == "__main__":
    create_docx()
