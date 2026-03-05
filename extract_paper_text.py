import docx
import os

def extract_text_from_docx(file_path, output_path):
    print(f"Extracting text from {file_path}...")
    try:
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        
        # Also extracting text from tables if any provided
        for table in doc.tables:
          text.append("\n--- Table Start ---\n")
          for row in table.rows:
              row_text = [cell.text for cell in row.cells]
              text.append(" | ".join(row_text))
          text.append("\n--- Table End ---\n")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(text))
        print(f"Successfully extracted to {output_path}")
    except Exception as e:
        print(f"Error extracting {file_path}: {e}")

base_dir = r"c:\Users\mitta\OneDrive\Documents\Major Project\Phase_2\nexus-llm-analytics-dist(Main)\nexus-llm-analytics-dist"
output_dir = r"c:\Users\mitta\.gemini\antigravity\brain\267ee7b6-76db-4cfc-9ad1-d69c31f04cda"

file1 = os.path.join(base_dir, "modified paper.docx")
file2 = os.path.join(base_dir, "original paper.docx")

out1 = os.path.join(output_dir, "modified_paper.md")
out2 = os.path.join(output_dir, "original_paper.md")

extract_text_from_docx(file1, out1)
extract_text_from_docx(file2, out2)
